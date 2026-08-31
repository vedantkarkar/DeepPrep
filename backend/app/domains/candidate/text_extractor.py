import io
import os
from typing import BinaryIO
import pypdf
import docx

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB

class TextExtractionError(Exception):
    """Base exception for resume text extraction."""
    pass

class UnsupportedFileTypeError(TextExtractionError):
    pass

class OversizedFileError(TextExtractionError):
    pass

class EmptyDocumentError(TextExtractionError):
    pass

class CorruptDocumentError(TextExtractionError):
    pass

class ResumeTextExtractor:
    """Extracts raw plain text from PDF, DOCX, and TXT files safely."""

    @classmethod
    def extract_from_bytes(cls, file_bytes: bytes, filename: str) -> str:
        if not file_bytes:
            raise EmptyDocumentError("Supplied resume file is empty (0 bytes).")

        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            raise OversizedFileError(
                f"File size ({len(file_bytes) / (1024*1024):.2f}MB) exceeds 5MB limit."
            )

        ext = os.path.splitext(filename)[1].lower().strip(".")
        if ext not in ("pdf", "docx", "txt"):
            raise UnsupportedFileTypeError(
                f"Unsupported file format '.{ext}'. Supported formats are: PDF, DOCX, TXT."
            )

        # Basic magic header sanity check
        if ext == "pdf":
            if not file_bytes.startswith(b"%PDF-"):
                raise CorruptDocumentError("Invalid PDF header. File is not a valid PDF document.")
            return cls._extract_pdf(io.BytesIO(file_bytes))
        elif ext == "docx":
            # DOCX is a ZIP archive starting with PK\x03\x04
            if not file_bytes.startswith(b"PK\x03\x04"):
                raise CorruptDocumentError("Invalid DOCX header. File is not a valid DOCX document.")
            return cls._extract_docx(io.BytesIO(file_bytes))
        elif ext == "txt":
            return cls._extract_txt(file_bytes)
        else:
            raise UnsupportedFileTypeError(f"Unsupported format '.{ext}'.")

    @classmethod
    def _extract_pdf(cls, stream: BinaryIO) -> str:
        try:
            reader = pypdf.PdfReader(stream)
            if len(reader.pages) == 0:
                raise EmptyDocumentError("PDF document has 0 pages.")

            extracted_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    extracted_text.append(text)

            full_text = "\n".join(extracted_text).strip()
            if not full_text:
                raise EmptyDocumentError("PDF document contains no readable text (may be image-only scan).")
            return full_text
        except EmptyDocumentError:
            raise
        except Exception as e:
            raise CorruptDocumentError(f"Failed to extract text from PDF: {str(e)}") from e

    @classmethod
    def _extract_docx(cls, stream: BinaryIO) -> str:
        try:
            doc = docx.Document(stream)
            extracted_text = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    extracted_text.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            extracted_text.append(cell.text.strip())

            full_text = "\n".join(extracted_text).strip()
            if not full_text:
                raise EmptyDocumentError("DOCX document contains no readable text.")
            return full_text
        except EmptyDocumentError:
            raise
        except Exception as e:
            raise CorruptDocumentError(f"Failed to extract text from DOCX: {str(e)}") from e

    @classmethod
    def _extract_txt(cls, file_bytes: bytes) -> str:
        try:
            text = file_bytes.decode("utf-8").strip()
            if not text:
                raise EmptyDocumentError("Text file is empty.")
            return text
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode("latin-1").strip()
                if not text:
                    raise EmptyDocumentError("Text file is empty.")
                return text
            except Exception as e:
                raise CorruptDocumentError(f"Failed to decode text file: {str(e)}") from e
